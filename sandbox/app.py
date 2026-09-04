"""Hardened document parser. Runs with no network egress and no API keys."""
from __future__ import annotations

import asyncio
import io
import os
import shutil
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from charset_normalizer import from_bytes
from fastapi import FastAPI, File, Query, UploadFile

MAX_BYTES = 512 * 1024 * 1024
TEXT_SUFFIXES = {".txt", ".md", ".rtf", ".csv", ".json", ".log", ".text", ""}


def _parse_pdf(data: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages), len(pages)


def _parse_docx(data: bytes) -> tuple[str, int]:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts), 1


def _parse_text(data: bytes) -> tuple[str, int]:
    best = from_bytes(data).best()
    return (str(best) if best else data.decode("utf-8", errors="replace")), 1


def _tesseract_version() -> str | None:
    if not shutil.which("tesseract"):
        return None
    try:
        proc = subprocess.run(["tesseract", "--version"], capture_output=True,
                              text=True, timeout=5, check=False)
    except (OSError, subprocess.TimeoutExpired):
        return None
    line = (proc.stdout or proc.stderr or "").splitlines()[:1]
    return line[0].strip() if line else None


def _ocr_workers() -> int:
    raw = os.environ.get("OCR_WORKERS", "").strip()
    if raw.isdigit() and int(raw) > 0:
        return int(raw)
    return max(1, os.cpu_count() or 1)


def _ocr_page(image: Path) -> str:
    proc = subprocess.run(
        ["tesseract", str(image), "stdout", "--oem", "1", "--psm", "6"],
        check=True, timeout=600, capture_output=True)
    return proc.stdout.decode("utf-8", errors="replace")


def _parse_ocr(data: bytes, suffix: str, workers: int | None = None) -> tuple[str, int]:
    if not shutil.which("tesseract"):
        raise RuntimeError("tesseract is not installed in this image")
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / f"input{suffix or '.pdf'}"
        src.write_bytes(data)
        if suffix.lower() == ".pdf":
            if not shutil.which("pdftoppm"):
                raise RuntimeError("pdftoppm is not installed in this image")
            subprocess.run(["pdftoppm", "-r", "200", "-png", str(src), f"{tmp}/page"],
                           check=True, timeout=600, capture_output=True)
            images = sorted(Path(tmp).glob("page*.png"))
        else:
            images = [src]
        planned = workers if workers and workers > 0 else _ocr_workers()
        workers = min(len(images), planned) if images else 1
        if workers <= 1:
            chunks = [_ocr_page(image) for image in images]
        else:
            with ThreadPoolExecutor(max_workers=workers) as pool:
                chunks = list(pool.map(_ocr_page, images))
        return "\n".join(chunks), len(images)


def create_app() -> FastAPI:
    app = FastAPI(title="Parser Sandbox", version="1.0.0")

    @app.get("/health")
    def health() -> dict:
        return {
            "status": "ok",
            "ocr_engine": "tesseract",
            "ocr_version": _tesseract_version(),
        }

    @app.post("/parse")
    async def parse(file: UploadFile = File(...), ocr: bool = False,
                    workers: int | None = Query(default=None)) -> dict:
        data = await file.read()
        if len(data) > MAX_BYTES:
            return {"text": "", "pages": 0, "parser": "none", "ok": False,
                    "reason": "file exceeds size limit"}

        suffix = Path(file.filename or "").suffix.lower()
        try:
            if ocr:
                text, pages = await asyncio.to_thread(_parse_ocr, data, suffix, workers)
                parser = "ocr"
            elif suffix == ".pdf":
                text, pages = await asyncio.to_thread(_parse_pdf, data)
                parser = "pypdf"
            elif suffix in {".docx", ".doc"}:
                text, pages = await asyncio.to_thread(_parse_docx, data)
                parser = "python-docx"
            elif suffix in TEXT_SUFFIXES:
                text, pages = _parse_text(data)
                parser = "text"
            else:
                return {"text": "", "pages": 0, "parser": "none", "ok": False,
                        "reason": f"unsupported extension {suffix!r}"}
        except Exception as exc:  # noqa: BLE001 - the sandbox reports, never raises
            return {"text": "", "pages": 0, "parser": "none", "ok": False,
                    "reason": f"{type(exc).__name__}: {exc}"}

        if not text.strip():
            return {"text": "", "pages": pages, "parser": parser, "ok": False,
                    "reason": "no extractable text"}
        return {"text": text, "pages": pages, "parser": parser, "ok": True, "reason": None}

    return app


app = create_app()
